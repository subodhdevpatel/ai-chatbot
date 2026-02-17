import os
import chromadb
from pypdf import PdfReader
from chromadb.utils import embedding_functions
from typing import List

# Configuration
DOCUMENTS_DIR = os.path.join(os.path.dirname(__file__), "documents")
DB_PATH = "chroma_db"
COLLECTION_NAME = "rag_collection"

def load_documents(directory: str) -> List[str]:
    """Reads all .txt files from the directory and returns their content."""
    documents = []
    if not os.path.exists(directory):
        os.makedirs(directory)
        return []
    
    for filename in os.listdir(directory):
        if filename.endswith(".txt"):
            path = os.path.join(directory, filename)
            with open(path, "r", encoding="utf-8") as f:
                documents.append(f.read())
            print(f"Loaded: {filename}")
        
        elif filename.endswith(".pdf"):
            try:
                reader = PdfReader(path)
                text = ""
                for page in reader.pages:
                    extract = page.extract_text()
                    if extract:
                        text += extract + "\n"
                documents.append(text)
                print(f"Loaded: {filename}")
            except Exception as e:
                print(f"Error loading {filename}: {e}")
                
        elif filename.endswith(".docx"):
            try:
                from docx import Document
                doc = Document(path)
                text = "\n".join([para.text for para in doc.paragraphs])
                documents.append(text)
                print(f"Loaded: {filename}")
            except Exception as e:
                print(f"Error loading {filename}: {e}")
                
    return documents

def split_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """Splits text into chunks with overlap."""
    chunks = []
    start = 0
    text_len = len(text)
    
    while start < text_len:
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start += (chunk_size - overlap)
        
    return chunks

def ingest_docs(file_path: str = None):
    """
    Reads documents, chunks them, and upserts them into ChromaDB.
    If file_path is provided, only ingests that specific file.
    Otherwise, reads all files in DOCUMENTS_DIR.
    Returns True if successful, False otherwise.
    """
    print("Initializing ChromaDB client...")
    try:
        client = chromadb.PersistentClient(path=DB_PATH)
        
        print("Initializing Embedding Function (all-MiniLM-L6-v2)...")
        # Uses the local model from sentence-transformers
        sentence_transformer_ef = embedding_functions.SentenceTransformerEmbeddingFunction(
            model_name="all-MiniLM-L6-v2"
        )
        
        collection = client.get_or_create_collection(
            name=COLLECTION_NAME,
            embedding_function=sentence_transformer_ef
        )
        
        raw_docs = []
        doc_names = []
        
        if file_path:
             print(f"Loading single document: {file_path}")
             # Re-use the logic inside load_documents but for a single file
             filename = os.path.basename(file_path)
             if filename.endswith(".txt"):
                with open(file_path, "r", encoding="utf-8") as f:
                    raw_docs.append(f.read())
                    doc_names.append(filename)
             elif filename.endswith(".pdf"):
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    extract = page.extract_text()
                    if extract:
                        text += extract + "\n"
                raw_docs.append(text)
                doc_names.append(filename)
             elif filename.endswith(".docx"):
                from docx import Document
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                raw_docs.append(text)
                doc_names.append(filename)
        else:
            print(f"Loading documents from '{DOCUMENTS_DIR}'...")
            raw_docs = load_documents(DOCUMENTS_DIR)
            # Note: load_documents doesn't return filenames, so we can't easily track IDs by filename here without refactoring load_documents too.
            # But for the bulk load it's fine.

        if not raw_docs:
            print("No documents found or extracted (raw_docs is empty).")
            return False

        # Filter out empty strings from raw_docs
        valid_docs = []
        valid_names = []
        for doc, name in zip(raw_docs, doc_names):
            if doc.strip():
                valid_docs.append(doc)
                valid_names.append(name)
            else:
                print(f"Warning: Document '{name}' was empty after extraction.")

        if not valid_docs:
             print("No valid text content found in any document.")
             return False

        raw_docs = valid_docs
        doc_names = valid_names

        all_chunks = []
        all_ids = []
        all_metadatas = []
        
        # We need unique IDs to avoid overwriting or duplicates if possible, 
        # but for now we'll stick to a simple counter or hash.
        # Ideally, we should delete old chunks for this file before re-ingesting.
        
        chunk_global_counter = 0 
        
        for i, doc in enumerate(raw_docs):
            chunks = split_text(doc)
            source_name = doc_names[i] if doc_names else "bulk_import"
            
            for chunk in chunks:
                all_chunks.append(chunk)
                # Make ID unique-ish
                import uuid
                all_ids.append(f"{source_name}_{uuid.uuid4().hex[:8]}")
                all_metadatas.append({"source": source_name})
                chunk_global_counter += 1
                
        if all_chunks:
            print(f"Upserting {len(all_chunks)} chunks into ChromaDB...")
            collection.upsert(
                documents=all_chunks,
                ids=all_ids,
                metadatas=all_metadatas
            )
            print("Ingestion complete!")
            return True
        else:
            print("No text content could be extracted.")
            return False

            
    except Exception as e:
        print(f"Error during ingestion: {e}")
        return False

def remove_document_from_db(filename: str):
    """
    Removes all chunks associated with a given filename from ChromaDB.
    """
    print(f"Removing document '{filename}' from ChromaDB...")
    try:
        client = chromadb.PersistentClient(path=DB_PATH)
        collection = client.get_collection(name=COLLECTION_NAME) # Don't create if not exists
        
        # Delete items where metadata 'source' is the filename
        collection.delete(where={"source": filename})
        print(f"Successfully removed '{filename}' from ChromaDB.")
        return True
    except Exception as e:
        print(f"Error removing document from ChromaDB: {e}")
        return False
        
def main():
    ingest_docs()

if __name__ == "__main__":
    main()
